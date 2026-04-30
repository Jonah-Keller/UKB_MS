"""Post-process rendered .docx for manuscript formatting.

Applies formatting that cannot be expressed in pandoc markdown:
- 1-inch margins on all sides
- Yellow highlight on ``[REF]`` placeholders
- Line numbers (restart each page)
- Page numbers (footer, right-aligned)

Originally authored for the NfL_Cardiovascular_UKB project (Keller et al.);
ported verbatim with minor generalisation for use across Ligh Lab projects.
"""

from __future__ import annotations

import copy
import logging
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

log = logging.getLogger(__name__)

_REF_RE = re.compile(r"\[REF\]")


def _get_run_text(r_elem) -> str:
    t_elem = r_elem.find(qn("w:t"))
    return t_elem.text if t_elem is not None and t_elem.text else ""


def _set_run_text(r_elem, text: str) -> None:
    t_elem = r_elem.find(qn("w:t"))
    if t_elem is None:
        t_elem = OxmlElement("w:t")
        r_elem.append(t_elem)
    t_elem.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t_elem.set(qn("xml:space"), "preserve")


def _make_highlight_yellow(r_elem) -> None:
    rPr = r_elem.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r_elem.insert(0, rPr)
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), "yellow")
    rPr.append(highlight)


def _split_run_on_pattern(r_elem, pattern, format_fn):
    text = _get_run_text(r_elem)
    if not text:
        return [r_elem]

    matches = list(pattern.finditer(text))
    if not matches:
        return [r_elem]

    result = []
    last_end = 0
    for m in matches:
        if m.start() > last_end:
            new_r = copy.deepcopy(r_elem)
            _set_run_text(new_r, text[last_end:m.start()])
            result.append(new_r)

        new_r = copy.deepcopy(r_elem)
        _set_run_text(new_r, m.group())
        format_fn(new_r)
        result.append(new_r)

        last_end = m.end()

    if last_end < len(text):
        new_r = copy.deepcopy(r_elem)
        _set_run_text(new_r, text[last_end:])
        result.append(new_r)

    return result


def _process_paragraph_runs(paragraph, pattern, format_fn) -> None:
    p_elem = paragraph._p
    run_elems = list(p_elem.findall(qn("w:r")))

    for r_elem in run_elems:
        new_elems = _split_run_on_pattern(r_elem, pattern, format_fn)
        if len(new_elems) == 1 and new_elems[0] is r_elem:
            continue
        parent = r_elem.getparent()
        idx = list(parent).index(r_elem)
        parent.remove(r_elem)
        for i, new_r in enumerate(new_elems):
            parent.insert(idx + i, new_r)


def _add_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            p = footer.paragraphs[0]
        else:
            p = footer.add_paragraph()

        p.alignment = 2  # WD_ALIGN_PARAGRAPH.RIGHT

        run = p.add_run()
        run.font.name = "Arial"
        run.font.size = Pt(10)

        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar_begin)

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run._r.append(instrText)

        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar_end)


def postprocess_docx(
    docx_path: Path,
    margin_in: float = 1.0,
    line_numbers: bool = True,
    highlight_ref_placeholder: bool = True,
    page_numbers: bool = True,
) -> None:
    """Apply manuscript formatting to a pandoc-generated .docx (modifies in place).

    Args:
        docx_path: Path to the .docx file (modified in place).
        margin_in: Page margin in inches (all sides). Default 1.0.
        line_numbers: If True, add per-page-restart line numbers.
        highlight_ref_placeholder: If True, highlight literal ``[REF]`` in yellow.
        page_numbers: If True, add right-aligned page numbers in footer.
    """
    doc = Document(str(docx_path))

    for section in doc.sections:
        section.top_margin    = Inches(margin_in)
        section.bottom_margin = Inches(margin_in)
        section.left_margin   = Inches(margin_in)
        section.right_margin  = Inches(margin_in)

    if highlight_ref_placeholder:
        for para in doc.paragraphs:
            if "[REF]" in para.text:
                _process_paragraph_runs(para, _REF_RE, _make_highlight_yellow)

    if line_numbers:
        for section in doc.sections:
            sectPr = section._sectPr
            lnNumType = OxmlElement("w:lnNumType")
            lnNumType.set(qn("w:countBy"), "1")
            lnNumType.set(qn("w:restart"), "newPage")
            sectPr.append(lnNumType)

    if page_numbers:
        _add_page_numbers(doc)

    doc.save(str(docx_path))
    log.info("Post-processed %s", docx_path.name)


if __name__ == "__main__":
    import sys
    logging.basicConfig(level=logging.INFO)
    if len(sys.argv) != 2:
        print("Usage: python -m analysis.helpers.docx_formatter <path.docx>")
        sys.exit(2)
    postprocess_docx(Path(sys.argv[1]))
