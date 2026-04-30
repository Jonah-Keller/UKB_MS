"""
docx_formatter.py — Post-process a manuscript .docx for journal submission.

Applies:
  1. Narrow margins (1 inch all sides)
  2. Continuous line numbers in the left margin
  3. Page numbers in the footer (centred, Arial 10 pt)
  4. Yellow highlight on any run containing "[REF]" placeholder text

Requires: python-docx  (pip install python-docx)
"""

from __future__ import annotations
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_COLOR_INDEX
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False


def _set_margins(doc: "Document", top: float = 1.0, bottom: float = 1.0,
                 left: float = 1.0, right: float = 1.0) -> None:
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)


def _enable_line_numbers(doc: "Document") -> None:
    """Add continuous line numbers to every section via sectPr/lnNumType."""
    for section in doc.sections:
        sectPr = section._sectPr
        lnNumType = OxmlElement("w:lnNumType")
        lnNumType.set(qn("w:countBy"), "1")
        lnNumType.set(qn("w:restart"), "continuous")
        lnNumType.set(qn("w:start"), "1")
        lnNumType.set(qn("w:distance"), "360")  # ~0.25 inch from text
        sectPr.append(lnNumType)


def _add_page_numbers(doc: "Document") -> None:
    """Add centred page number to footer of each section."""
    for section in doc.sections:
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        para = footer.paragraphs[0]
        para.clear()
        para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.font.name = "Arial"
        run.font.size = Pt(10)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)


def _highlight_ref_placeholders(doc: "Document") -> None:
    """Yellow-highlight any run containing '[REF]'."""
    for para in doc.paragraphs:
        for run in para.runs:
            if "[REF]" in run.text:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _add_running_header(doc: "Document", header_text: str) -> None:
    """Add a centred running header to every section."""
    for section in doc.sections:
        header = section.header
        if not header.paragraphs:
            header.add_paragraph()
        para = header.paragraphs[0]
        para.clear()
        para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header_text)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.italic = True


def _set_body_font(doc: "Document", font_name: str = "Arial",
                   font_size_pt: float = 12) -> None:
    """Set Normal style to the specified font and size; also apply to body runs."""
    from docx.shared import Pt
    style = doc.styles["Normal"]
    style.font.name  = font_name
    style.font.size  = Pt(font_size_pt)
    # Propagate to all paragraph runs so existing content picks up the change
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name is None:
                run.font.name = font_name
            if run.font.size is None:
                run.font.size = Pt(font_size_pt)


def postprocess_docx(path: Path, running_header: str = "") -> None:
    if not _HAS_DOCX:
        import warnings
        warnings.warn(
            "python-docx not installed — skipping docx post-processing. "
            "Install with: pip install python-docx",
            stacklevel=2,
        )
        return

    doc = Document(str(path))
    _set_margins(doc)
    _set_body_font(doc, font_name="Arial", font_size_pt=12)
    _enable_line_numbers(doc)
    _add_page_numbers(doc)
    _highlight_ref_placeholders(doc)
    if running_header:
        _add_running_header(doc, running_header)
    doc.save(str(path))
